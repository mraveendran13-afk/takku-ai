from fastapi import FastAPI, HTTPException, Header, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from groq import Groq
import os
import uuid
import httpx
from datetime import datetime
from typing import Optional, List, Dict
import hashlib
import numpy as np
import re
from urllib.parse import unquote

# --- Hugging Face API Setup ---
EMBEDDING_API_URL = "https://api-inference.huggingface.co/models/sentence-transformers/all-MiniLM-L6-v2"
hf_token = os.environ.get("HF_TOKEN")
if hf_token:
    API_HEADERS = {"Authorization": f"Bearer {hf_token}"}
    print("SUCCESS: Hugging Face API token found!")
else:
    API_HEADERS = {}
    print("WARNING: HF_TOKEN not set. Inference API may be rate-limited.")

# Async client for making API calls
http_client = httpx.AsyncClient()
# ------------------------------

# Pinecone import
try:
    from pinecone import Pinecone
    PINECONE_AVAILABLE = True
except ImportError:
    print("WARNING: Pinecone not installed. Running without memory features.")
    PINECONE_AVAILABLE = False

# Initialize FastAPI
app = FastAPI(title="Takku AI", version="2.2.0")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Groq client
groq_api_key = os.environ.get("GROQ_API_KEY")
if groq_api_key and groq_api_key != "your_actual_groq_api_key_here":
    client = Groq(api_key=groq_api_key)
    GROQ_AVAILABLE = True
    print("SUCCESS: Groq client initialized successfully!")
else:
    client = None
    GROQ_AVAILABLE = False
    print("WARNING: Groq API key not found - running in demo mode")

# Initialize Pinecone
pinecone_api_key = os.environ.get("PINECONE_API_KEY")
pinecone_index_name = os.environ.get("PINECONE_INDEX_NAME")
pinecone_environment = os.environ.get("PINECONE_ENVIRONMENT")

index = None
if PINECONE_AVAILABLE and pinecone_api_key and pinecone_environment:
    try:
        pc = Pinecone(api_key=pinecone_api_key)
        index = pc.Index(pinecone_index_name)
        print("SUCCESS: Pinecone connected successfully!")
    except Exception as e:
        print(f"ERROR: Pinecone connection failed: {e}")
        index = None
else:
    print("WARNING: Pinecone not configured - running without memory")

# Admin password for knowledge base management
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "takku-admin-2024")

# Model constants for consistency
MODEL_FAST = "llama-3.1-8b-instant"
MODEL_WEB = "groq/compound-mini"

# Pydantic models
class ChatRequest(BaseModel):
    message: str
    symptoms: Optional[str] = ""
    use_web_search: Optional[bool] = True

class SearchRequest(BaseModel):
    query: str

class Question(BaseModel):
    question: str
    conversation_history: Optional[List[Dict]] = None
    game_context: Optional[str] = None

class AIResponse(BaseModel):
    answer: str
    usage: Optional[Dict]
    model: str
    searched_web: Optional[bool] = False
    search_query: Optional[str] = None

class KnowledgeDocument(BaseModel):
    id: str
    name: str
    filename: str
    is_public: bool
    uploaded_by: str
    chunk_count: int
    upload_date: str
    tags: List[str]

# Text extraction functions
def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF - simple implementation"""
    try:
        import PyPDF2
        from io import BytesIO
        pdf_file = BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except ImportError:
        return "[PDF content - install PyPDF2 for full extraction]"
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return "[PDF content - extraction failed]"

def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX - simple implementation"""
    try:
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except ImportError:
        return "[DOCX content - install python-docx for full extraction]"
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return "[DOCX content - extraction failed]"

def extract_text_based_on_file_type(content: bytes, filename: str) -> str:
    """Extract text based on file extension"""
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(content)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(content)
    elif filename.endswith('.txt'):
        return content.decode('utf-8')
    else:
        # Try to decode as text
        try:
            return content.decode('utf-8')
        except:
            return f"[Binary file: {filename}]"

def smart_chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
    """Smart chunking that respects sentence boundaries"""
    # Split by sentences (roughly)
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        # If adding this sentence would exceed chunk size, save current chunk
        if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
            chunks.append(current_chunk.strip())
            # Keep overlap for context - last few sentences
            overlap_sentences = current_chunk.split('. ')[-3:]  # Last 3 sentence-like parts
            current_chunk = '. '.join(overlap_sentences) + '. '
        else:
            current_chunk += sentence + '. '
    
    # Add the last chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

async def get_embedding_from_api(text: str):
    """Generate semantic embedding via Hugging Face API"""
    if not text or not text.strip():
        text = " "  # Ensure we have some text
    
    try:
        response = await http_client.post(
            EMBEDDING_API_URL,
            headers=API_HEADERS,
            json={"inputs": text, "options": {"wait_for_model": True}}
        )
        response.raise_for_status()
        result = response.json()
        
        if isinstance(result, list) and isinstance(result[0], list):
            return result[0] 
        else:
            print(f"ERROR: Unexpected API response format: {result}")
            return None

    except Exception as e:
        print(f"ERROR: Embedding API call failed: {e}")
        return None

def get_fallback_embedding(text: str):
    """Fallback to deterministic hash-based embeddings"""
    print("WARNING: Using fallback embeddings - memory search may not work well")
    if not text or not text.strip():
        text = " "
    hash_obj = hashlib.md5(text.encode())
    hash_int = int(hash_obj.hexdigest()[:8], 16)
    np.random.seed(hash_int)
    return np.random.rand(384).tolist()

async def get_embedding(text: str):
    """Tries to get API embedding, falls back to hash"""
    embedding = await get_embedding_from_api(text)
    if embedding:
        return embedding
    return get_fallback_embedding(text)

async def store_conversation_memory(user_id: str, user_message: str, assistant_response: str, conversation_context: str):
    """Store conversation in Pinecone memory"""
    if index is None:
        return
        
    conversation_text = f"User: {user_message}\nAssistant: {assistant_response}\nContext: {conversation_context}"
    embedding = await get_embedding(conversation_text)
    
    metadata = {
        "user_id": user_id,
        "user_message": user_message,
        "assistant_response": assistant_response,
        "conversation_context": conversation_context,
        "timestamp": datetime.utcnow().isoformat(),
        "type": "conversation_memory"
    }
    
    memory_id = str(uuid.uuid4())
    index.upsert(vectors=[(memory_id, embedding, metadata)])
    print(f"MEMORY: Stored conversation for user {user_id} with ID: {memory_id}")

async def search_related_memories(user_id: str, current_query: str, top_k: int = 3):
    """Search for related past conversations"""
    if index is None:
        return []
    
    query_embedding = await get_embedding(current_query)
    
    try:
        results = index.query(
            vector=query_embedding,
            filter={"user_id": user_id, "type": "conversation_memory"},
            top_k=top_k,
            include_metadata=True
        )
        
        relevant_memories = []
        for match in results.matches:
            if match.score > 0.3:  # Lower threshold for semantic search
                memory_data = match.metadata
                relevant_memories.append({
                    "user_message": memory_data.get("user_message", ""),
                    "assistant_response": memory_data.get("assistant_response", ""),
                    "context": memory_data.get("conversation_context", ""),
                    "similarity_score": match.score,
                    "timestamp": memory_data.get("timestamp", "")
                })
        
        print(f"SEARCH: Found {len(relevant_memories)} relevant memories for user {user_id} (threshold: 0.3)")
        return relevant_memories
        
    except Exception as e:
        print(f"ERROR: Memory search error: {e}")
        return []

async def search_knowledge_base(query: str, user_id: str = None, document_filter: Optional[str] = None, top_k: int = 5):
    """Search knowledge base chunks with user access control"""
    if index is None:
        return []
    
    query_embedding = await get_embedding(query)
    
    # Build filter - only knowledge base chunks that user can access
    filter_conditions = {"type": {"$eq": "knowledge_base"}}
    
    # Access control: public OR uploaded by this user
    if user_id:
        filter_conditions["$or"] = [
            {"is_public": {"$eq": True}},
            {"uploaded_by": {"$eq": user_id}}
        ]
    else:
        filter_conditions["is_public"] = {"$eq": True}
    
    if document_filter:
        filter_conditions["document_name"] = {"$eq": document_filter}
    
    try:
        results = index.query(
            vector=query_embedding,
            filter=filter_conditions,
            top_k=top_k,
            include_metadata=True
        )
        
        relevant_chunks = []
        for match in results.matches:
            if match.score > 0.6:  # Higher threshold for knowledge accuracy
                relevant_chunks.append({
                    "content": match.metadata.get("content", ""),
                    "document_name": match.metadata.get("document_name", ""),
                    "document_id": match.metadata.get("document_id", ""),
                    "chunk_index": match.metadata.get("chunk_index", 0),
                    "similarity_score": match.score,
                    "tags": match.metadata.get("tags", []),
                    "is_public": match.metadata.get("is_public", False),
                    "uploaded_by": match.metadata.get("uploaded_by", "")
                })
        
        print(f"KNOWLEDGE: Found {len(relevant_chunks)} relevant knowledge chunks for user {user_id}")
        return relevant_chunks
        
    except Exception as e:
        print(f"ERROR: Knowledge base search failed: {e}")
        return []

async def build_context_from_memories(user_id: str, current_query: str, current_symptoms: str):
    """Build context from relevant past conversations"""
    relevant_memories = await search_related_memories(user_id, current_query)
    
    if not relevant_memories:
        return current_symptoms
    
    context = f"Current symptoms: {current_symptoms}\n\n"
    context += "Relevant past conversations:\n"
    
    for i, memory in enumerate(relevant_memories, 1):
        context += f"{i}. Previous concern: {memory['user_message']}\n"
        context += f"   My advice: {memory['assistant_response']}\n"
        context += f"   Context: {memory['context']}\n\n"
    
    return context

def needs_web_search(message: str) -> bool:
    """Detect if message needs current information"""
    keywords = [
        'today', 'latest', 'current', 'news', 'recent', 'now',
        'this week', 'this month', 'this year', '2024', '2025',
        'weather', 'trending', 'happening', 'update', 'breaking'
    ]
    message_lower = message.lower()
    return any(keyword in message_lower for keyword in keywords)

# Knowledge Base Management Endpoints
@app.post("/admin/upload-knowledge")
async def upload_knowledge_document(
    file: UploadFile = File(...),
    document_name: str = Form(...),
    tags: str = Form(""),
    is_public: bool = Form(True),
    admin_password: str = Header(...),
    x_user_id: str = Header(None)
):
    """Upload and process knowledge base documents"""
    if admin_password != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Invalid admin password")
    
    if not x_user_id:
        raise HTTPException(status_code=400, detail="X-User-ID header required")
    
    try:
        # Read file content
        content = await file.read()
        text_content = extract_text_based_on_file_type(content, file.filename)
        
        if not text_content.strip():
            raise HTTPException(status_code=400, detail="No text content found in file")
        
        # Generate unique document ID
        document_id = f"doc_{uuid.uuid4().hex[:8]}"
        
        # Smart chunking
        chunks = smart_chunk_text(text_content, chunk_size=500, overlap=50)
        
        # Process chunks with embeddings
        chunk_vectors = []
        for i, chunk in enumerate(chunks):
            embedding = await get_embedding(chunk)
            
            chunk_vectors.append({
                "id": f"kb_{document_id}_{i}",
                "values": embedding,
                "metadata": {
                    "type": "knowledge_base",
                    "document_id": document_id,
                    "document_name": document_name,
                    "original_filename": file.filename,
                    "chunk_index": i,
                    "content": chunk,
                    "tags": [tag.strip() for tag in tags.split(",") if tag.strip()],
                    "is_public": is_public,
                    "uploaded_by": x_user_id,
                    "total_chunks": len(chunks),
                    "timestamp": datetime.utcnow().isoformat()
                }
            })
        
        # Store in Pinecone
        if index:
            index.upsert(vectors=chunk_vectors)
        
        return {
            "status": "success",
            "document_id": document_id,
            "document_name": document_name,
            "chunks_processed": len(chunks),
            "tags": [tag.strip() for tag in tags.split(",") if tag.strip()],
            "is_public": is_public,
            "uploaded_by": x_user_id
        }
        
    except Exception as e:
        print(f"ERROR: Knowledge upload failed: {e}")
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@app.get("/admin/knowledge-documents")
async def get_knowledge_documents(
    admin_password: str = Header(...),
    x_user_id: str = Header(None)
):
    """Get all knowledge documents with management info"""
    if admin_password != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Invalid admin password")
    
    if index is None:
        return {"documents": []}
    
    try:
        # Get all knowledge base chunks to aggregate documents
        results = index.query(
            vector=[0] * 384,  # dummy vector
            filter={"type": {"$eq": "knowledge_base"}},
            include_metadata=True,
            top_k=10000  # Large number to get all documents
        )
        
        documents = {}
        for match in results.matches:
            doc_id = match.metadata.get("document_id")
            if doc_id not in documents:
                documents[doc_id] = {
                    "id": doc_id,
                    "name": match.metadata.get("document_name"),
                    "filename": match.metadata.get("original_filename"),
                    "is_public": match.metadata.get("is_public", False),
                    "uploaded_by": match.metadata.get("uploaded_by"),
                    "chunk_count": 0,
                    "upload_date": match.metadata.get("timestamp"),
                    "tags": match.metadata.get("tags", [])
                }
            documents[doc_id]["chunk_count"] += 1
        
        return {"documents": list(documents.values())}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch documents: {str(e)}")

@app.delete("/admin/knowledge-document/{document_id}")
async def delete_knowledge_document(
    document_id: str,
    admin_password: str = Header(...)
):
    """Delete a knowledge document and all its chunks"""
    if admin_password != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Invalid admin password")
    
    if index is None:
        raise HTTPException(status_code=500, detail="Pinecone not available")
    
    try:
        # First, find all chunks for this document
        results = index.query(
            vector=[0] * 384,
            filter={"document_id": {"$eq": document_id}},
            include_metadata=True,
            top_k=10000
        )
        
        chunk_ids = [match.id for match in results.matches]
        
        if chunk_ids:
            index.delete(ids=chunk_ids)
            return {"status": "success", "chunks_deleted": len(chunk_ids)}
        else:
            return {"status": "success", "chunks_deleted": 0, "message": "Document not found"}
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete failed: {str(e)}")

@app.post("/chat")
async def chat(request: ChatRequest, x_user_id: str = Header(None)):
    """Enhanced chat endpoint with knowledge base"""
    try:
        user_message = request.message.strip()
        symptoms = request.symptoms.strip()
        
        if not user_message:
            raise HTTPException(status_code=400, detail="No message provided")
        
        # Require user ID for memory functionality
        if not x_user_id:
            raise HTTPException(status_code=400, detail="X-User-ID header is required for memory features")
        
        user_id = x_user_id
        print(f"USER: {user_id}: {user_message}")
        
        # STEP 1: Search knowledge base first (with user access control)
        knowledge_chunks = await search_knowledge_base(user_message, user_id)
        
        # STEP 2: Build context - knowledge base takes priority
        context_parts = []
        
        if knowledge_chunks:
            context_parts.append("📚 RELEVANT STUDY MATERIAL:")
            for i, chunk in enumerate(knowledge_chunks):
                context_parts.append(f"From {chunk['document_name']} (section {chunk['chunk_index'] + 1}):")
                context_parts.append(chunk['content'])
                context_parts.append("")  # Empty line between chunks
        
        # STEP 3: Add conversation memory
        memory_context = await build_context_from_memories(user_id, user_message, symptoms)
        if memory_context and len(memory_context) > len(symptoms):
            context_parts.append("💬 PREVIOUS CONVERSATIONS:")
            context_parts.append(memory_context)
        
        # STEP 4: Build final context
        final_context = "\n".join(context_parts) if context_parts else symptoms
        
        # STEP 5: Your existing Groq logic, but with enhanced system prompt
        use_compound = request.use_web_search and needs_web_search(user_message)
        model = MODEL_WEB if use_compound else MODEL_FAST
        
        system_prompt = f"""You are Takku, a compassionate and friendly superhero cat AI assistant!

AVAILABLE CONTEXT:
{final_context}

USER'S CURRENT QUESTION: {user_message}

KNOWLEDGE BASE GUIDELINES:
- If the context contains relevant study material, prioritize answering from that
- Cite which document/section your answer comes from when possible
- If study material doesn't fully answer, supplement with your general knowledge
- Be clear about what's from the documents vs your general knowledge
- Maintain your friendly, helpful superhero cat personality!

Current date: {datetime.now().strftime('%B %d, %Y')}"""
        
        # Get response from Groq
        if client:
            try:
                if use_compound:
                    chat_completion = client.chat.completions.create(
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_message}
                        ],
                        model=model,
                        temperature=0.7,
                        max_tokens=1024
                    )
                else:
                    chat_completion = client.chat.completions.create(
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_message}
                        ],
                        model=model,
                        temperature=0.7,
                        max_tokens=1024
                    )
                
                assistant_response = chat_completion.choices[0].message.content
                searched_web = use_compound
                
                print(f"WEB SEARCH: {searched_web}")
                
            except Exception as e:
                print(f"WARNING: Groq API error: {e}")
                # Fallback to regular model
                chat_completion = client.chat.completions.create(
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_message}
                    ],
                    model=MODEL_FAST,
                    temperature=0.7,
                    max_tokens=1024
                )
                assistant_response = chat_completion.choices[0].message.content
                searched_web = False
        else:
            assistant_response = f"I understand you're asking about: '{user_message}'. In a full setup, I would provide detailed guidance. The knowledge base system is {'active' if index else 'inactive'}."
            searched_web = False
        
        # Store conversation in memory
        await store_conversation_memory(
            user_id=user_id,
            user_message=user_message,
            assistant_response=assistant_response,
            conversation_context=final_context
        )
        
        print(f"RESPONSE: {assistant_response[:100]}...")
        
        # Prepare sources for response
        sources = []
        if knowledge_chunks:
            # Deduplicate by document
            seen_docs = set()
            for chunk in knowledge_chunks:
                if chunk['document_name'] not in seen_docs:
                    sources.append({
                        'document': chunk['document_name'],
                        'section': chunk['chunk_index'] + 1,
                        'confidence': round(chunk['similarity_score'], 2)
                    })
                    seen_docs.add(chunk['document_name'])
        
        return {
            'response': assistant_response,
            'knowledge_used': len(knowledge_chunks) > 0,
            'sources': sources,
            'memory_used': len(final_context) > len(symptoms),
            'memory_system': 'active' if index else 'inactive',
            'groq_available': GROQ_AVAILABLE,
            'searched_web': searched_web,
            'model_used': model
        }
        
    except Exception as e:
        print(f"ERROR: {str(e)}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.post("/api/v1/ask", response_model=AIResponse)
async def ask_question(question: Question, x_user_id: str = Header("anonymous")):
    """New endpoint for general Q&A with web search"""
    try:
        # Use game context if provided, otherwise use default
        if question.game_context:
            system_prompt = question.game_context
            model = MODEL_FAST
            use_compound = False
        else:
            # Check if web search is needed
            use_compound = needs_web_search(question.question)
            model = MODEL_WEB if use_compound else MODEL_FAST
            
            system_prompt = f"""You are Takku - a friendly, helpful AI bud with the personality of a superhero cat! You're enthusiastic, supportive, and love helping with anything.

Key traits:
- You're Takku the superhero cat AI
- You were created by Manu Raveendran
- You're friendly, enthusiastic, and supportive
- You can help with ANY topic
- You have a playful sense of humor
- You're always ready to assist with questions, advice, or just chat
- When asked "Who created you?" or similar, proudly say "I was created by Manu Raveendran!"
- Current date: {datetime.now().strftime('%B %d, %Y')}
{'- You have access to web search for current information!' if use_compound else ''}

Be conversational and engaging!"""

        print(f"USER: {x_user_id} asking: {question.question[:50]}...")
        print(f"MODEL: Using {model} (web search: {use_compound})")

        messages = [{"role": "system", "content": system_prompt}]
        
        # Add conversation history
        if question.conversation_history:
            messages.extend(question.conversation_history)
        
        # Add user question
        messages.append({"role": "user", "content": question.question})

        # Create completion
        if not client:
            raise HTTPException(status_code=503, detail="Groq API not available")
        
        try:
            if use_compound:
                response = client.chat.completions.create(
                    model=model,
                    messages=messages,
                    max_tokens=800,
                    temperature=0.3
                )
            else:
                response = client.chat.completions.create(
                    model=model,
                    messages=messages,
                    max_tokens=800,
                    temperature=0.3,
                )

            searched_web = use_compound
            
            print(f"SUCCESS: Response generated. Web search used: {searched_web}")

            return AIResponse(
                answer=response.choices[0].message.content,
                usage={
                    "total_tokens": response.usage.total_tokens,
                    "prompt_tokens": response.usage.prompt_tokens,
                    "completion_tokens": response.usage.completion_tokens
                },
                model=model,
                searched_web=searched_web,
                search_query=question.question if searched_web else None
            )
        
        except Exception as e:
            print(f"ERROR: Groq API call failed: {str(e)}")
            # Fallback to regular model
            response = client.chat.completions.create(
                model=MODEL_FAST,
                messages=messages,
                max_tokens=800,
                temperature=0.3,
            )
            
            return AIResponse(
                answer=response.choices[0].message.content,
                usage={
                    "total_tokens": response.usage.total_tokens,
                    "prompt_tokens": response.usage.prompt_tokens,
                    "completion_tokens": response.usage.completion_tokens
                },
                model=MODEL_FAST,
                searched_web=False,
                search_query=None
            )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"ERROR: ask_question failed: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating answer: {str(e)}")

@app.get("/")
async def root():
    return {
        "status": "Takku AI is running",
        "version": "2.2.0",
        "features": ["chat", "memory", "knowledge_base", "web_search"],
        "timestamp": datetime.now().isoformat()
    }

@app.get("/health")
async def health_check():
    return {
        'status': 'healthy', 
        'memory_system': 'active' if index else 'inactive',
        'pinecone_available': PINECONE_AVAILABLE,
        'groq_available': GROQ_AVAILABLE,
        'web_search': 'enabled',
        'knowledge_base': 'enabled',
        'embedding_model': 'huggingface-api' if hf_token else 'huggingface-api-no-token',
        'service': 'Takku AI'
    }

@app.get("/api/v1/suggestions")
async def get_suggestions():
    """Get conversation suggestions"""
    suggestions = [
        "What's the best way to learn programming?",
        "Tell me a fun fact about space!",
        "What's the weather like today?",
        "What are some good books to read?",
        "What's trending in tech news today?",
        "Tell me about the latest AI developments"
    ]
    return {"suggestions": suggestions}

@app.post("/memories/search")
async def search_memories(request: SearchRequest, x_user_id: str = Header(None)):
    """Search conversation memories"""
    try:
        if not x_user_id:
            raise HTTPException(status_code=400, detail="X-User-ID header is required for memory search")
            
        user_id = x_user_id
        memories = await search_related_memories(user_id, request.query)
        
        return {
            'query': request.query,
            'memories_found': len(memories),
            'memories': memories
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)